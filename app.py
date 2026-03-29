import os
import re
import uuid
import sqlite3
from datetime import datetime
from typing import List, Dict, Any, Optional
from io import BytesIO

import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from dotenv import load_dotenv
from PIL import Image
from google import genai

load_dotenv()

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="ABT DATA ANALYST REPORTS",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================
# CONSTANTS
# =========================================================
APP_NAME = "ABT DATA ANALYST REPORTS"
LOGO_PATH = "ABT LOGO.jpg"
DB_PATH = "abt_data_analyst_history.db"
DEFAULT_MODEL = "gemini-2.5-flash"

# =========================================================
# CSS
# =========================================================
st.markdown(
    """
<style>
:root {
    --bg: #06080d;
    --panel: #0d1118;
    --panel2: #121827;
    --line: rgba(255,255,255,0.08);
    --text: #f3f6ff;
    --muted: #99a4bf;
    --blue: #2563eb;
    --red: #ef4444;
    --cyan: #06b6d4;
    --purple: #7c3aed;
}

html, body, [class*="css"] {
    font-family: "Inter", sans-serif;
}

.stApp {
    background:
      radial-gradient(circle at top left, rgba(37,99,235,0.18), transparent 24%),
      radial-gradient(circle at top right, rgba(239,68,68,0.15), transparent 20%),
      radial-gradient(circle at bottom center, rgba(124,58,237,0.12), transparent 24%),
      linear-gradient(180deg, #05070b 0%, #090d14 100%);
    color: var(--text);
}

.block-container {
    max-width: 1400px;
    padding-top: 1rem;
    padding-bottom: 2rem;
}

.shell {
    background: linear-gradient(180deg, rgba(255,255,255,0.02), rgba(255,255,255,0.01));
    border: 1px solid var(--line);
    border-radius: 28px;
    padding: 22px;
    box-shadow: 0 24px 80px rgba(0,0,0,0.35);
}

.hero {
    border-radius: 26px;
    padding: 24px;
    background:
      linear-gradient(135deg, rgba(37,99,235,0.17), rgba(239,68,68,0.10), rgba(124,58,237,0.10)),
      linear-gradient(180deg, #0b1018 0%, #0a0f16 100%);
    border: 1px solid var(--line);
    margin-bottom: 18px;
}

.hero-title {
    font-size: 2.15rem;
    font-weight: 900;
    letter-spacing: -0.02em;
    color: white;
    margin-bottom: .2rem;
}

.hero-sub {
    color: #c4cee6;
    font-size: 1.02rem;
}

.badge-row {
    display: flex;
    gap: 10px;
    flex-wrap: wrap;
    margin-top: 14px;
}

.badge-pill {
    background: rgba(255,255,255,0.05);
    border: 1px solid var(--line);
    color: #dbe3f5;
    padding: 8px 12px;
    border-radius: 999px;
    font-size: .86rem;
}

.section-card {
    background: linear-gradient(180deg, rgba(17,22,33,0.97), rgba(10,14,21,0.98));
    border: 1px solid var(--line);
    border-radius: 22px;
    padding: 18px;
    margin-bottom: 16px;
}

.info-box {
    background: linear-gradient(135deg, rgba(37,99,235,0.14), rgba(6,182,212,0.10));
    border-left: 4px solid #38bdf8;
    border: 1px solid var(--line);
    border-radius: 16px;
    padding: 14px 16px;
    color: #ebf2ff;
    margin: 10px 0;
}

.stat-card {
    background: linear-gradient(180deg, #111723, #0d1320);
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 14px 16px;
}

.stat-label {
    color: var(--muted);
    font-size: .84rem;
}

.stat-value {
    color: #ffffff;
    font-size: 1.35rem;
    font-weight: 800;
}

.history-card {
    background: linear-gradient(180deg, #101521, #0c1018);
    border: 1px solid var(--line);
    border-radius: 16px;
    padding: 12px 14px;
    margin-bottom: 10px;
}

.stTextArea textarea, .stTextInput input {
    background: #0b1018 !important;
    color: #f8fafc !important;
    border-radius: 14px !important;
}

.stButton > button {
    border-radius: 14px !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
    background: linear-gradient(135deg, #2563eb, #ef4444) !important;
    color: white !important;
    font-weight: 800 !important;
    padding: 0.6rem 1rem !important;
}

.stDownloadButton > button {
    border-radius: 14px !important;
    font-weight: 800 !important;
}

[data-testid="stSidebar"] {
    background: linear-gradient(180deg, rgba(11,14,20,0.98), rgba(8,10,15,0.98));
    border-right: 1px solid var(--line);
}
</style>
""",
    unsafe_allow_html=True,
)

# =========================================================
# DATABASE
# =========================================================
def init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS report_history (
            id TEXT PRIMARY KEY,
            created_at TEXT,
            title TEXT,
            content_type TEXT,
            content TEXT
        )
        """
    )
    conn.commit()
    conn.close()


def save_history(title: str, content_type: str, content: str) -> None:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO report_history (id, created_at, title, content_type, content) VALUES (?, ?, ?, ?, ?)",
        (
            str(uuid.uuid4()),
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            title,
            content_type,
            content,
        ),
    )
    conn.commit()
    conn.close()


def load_history(limit: int = 100) -> List[Dict[str, Any]]:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, created_at, title, content_type, content
        FROM report_history
        ORDER BY created_at DESC
        LIMIT ?
        """,
        (limit,),
    )
    rows = cur.fetchall()
    conn.close()
    return [
        {
            "id": row[0],
            "created_at": row[1],
            "title": row[2],
            "content_type": row[3],
            "content": row[4],
        }
        for row in rows
    ]


def delete_history(item_id: str) -> None:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM report_history WHERE id = ?", (item_id,))
    conn.commit()
    conn.close()


# =========================================================
# SESSION STATE
# =========================================================
defaults = {
    "raw_analysis_text": "",
    "empirical_review_text": "",
    "study_context_text": "",
    "parsed_tables": [],
    "interpreted_tables": [],
    "summary_of_findings": "",
    "discussion_of_findings": "",
    "final_outputs": "",
    "full_report": "",
    "comprehensive_report": "",
    "limitations_future_research": "",
    "pdf_buffer": None,
    "docx_buffer": None,
    "history_name_input": "",
    "opened_history_title": "",
    "opened_history_content": "",
}
for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val


# =========================================================
# API
# =========================================================
def get_api_key() -> str:
    key = os.getenv("GEMINI_API_KEY", "")
    if key:
        return key.strip()

    try:
        key = st.secrets.get("GEMINI_API_KEY", "")
        if key:
            return key.strip()
    except Exception:
        pass

    return ""


def get_client(api_key: str):
    return genai.Client(api_key=api_key)


def gemini_text(api_key: str, prompt: str, model_name: str = DEFAULT_MODEL) -> str:
    try:
        client = get_client(api_key)
        response = client.models.generate_content(model=model_name, contents=prompt)
        return (response.text or "").strip()
    except Exception as e:
        return f"ERROR: {e}"


# =========================================================
# TEXT HELPERS
# =========================================================
def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = text.replace("\uf0b7", "•")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(uploaded_file) -> str:
    try:
        uploaded_file.seek(0)
        reader = PdfReader(uploaded_file)
        pages = [page.extract_text() or "" for page in reader.pages]
        return clean_text("\n".join(pages))
    except Exception as e:
        return f"ERROR: Could not read PDF. {e}"


def extract_docx_text(uploaded_file) -> str:
    try:
        uploaded_file.seek(0)
        doc = Document(uploaded_file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                rows.append("\t".join(cells))
            table_texts.append("\n".join(rows))

        return clean_text("\n\n".join(paragraphs + table_texts))
    except Exception as e:
        return f"ERROR: Could not read DOCX. {e}"


def extract_text_from_upload(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)
    if name.endswith(".docx"):
        return extract_docx_text(uploaded_file)
    return "ERROR: Unsupported file type."


# =========================================================
# PARSING HELPERS
# =========================================================
WORD_TO_NUM = {
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
}

TABLE_START_PATTERN = re.compile(
    r"^\s*table\s+(\d+)\s*[:.\-]?\s*(.*)$",
    flags=re.I,
)


def to_number(token: str) -> Optional[int]:
    token = str(token).strip().lower()
    if token.isdigit():
        return int(token)
    return WORD_TO_NUM.get(token)


def extract_table_number(line: str) -> Optional[int]:
    match = TABLE_START_PATTERN.match(str(line).strip())
    if match:
        return int(match.group(1))
    return None


def extract_table_title_line(text: str) -> str:
    for line in str(text).split("\n"):
        if extract_table_number(line) is not None:
            return line.strip()
    return "Untitled Table"


def table_number_from_text(text: str) -> int:
    for line in str(text).split("\n"):
        num = extract_table_number(line)
        if num is not None:
            return num
    return 999


def section_sort_number(tag: str) -> int:
    tag = str(tag)
    match = re.search(r"\b(\d+)\b", tag)
    if match:
        return int(match.group(1))

    word_match = re.search(
        r"\b(one|two|three|four|five|six|seven|eight|nine|ten)\b",
        tag.lower(),
    )
    if word_match:
        return WORD_TO_NUM[word_match.group(1)]
    return 999


def split_tables_from_text(raw_text: str) -> List[str]:
    raw_text = clean_text(raw_text)
    if not raw_text:
        return []

    lines = [ln.rstrip() for ln in raw_text.split("\n") if ln.strip()]
    if not lines:
        return []

    tables: List[str] = []
    current_block: List[str] = []
    pending_context: List[str] = []

    def is_table_start(line: str) -> bool:
        return extract_table_number(line) is not None

    def is_context_heading(line: str) -> bool:
        lowered = line.strip().lower()
        patterns = [
            r"^answering research questions$",
            r"^hypotheses testing$",
            r"^research question\s+(one|two|three|four|five|six|\d+)",
            r"^question\s+(one|two|three|four|five|six|\d+)",
            r"^rq\s*(one|two|three|four|five|six|\d+)",
            r"^h0?\d+",
            r"^hypothesis\s+(one|two|three|four|five|six|\d+)",
            r"^research hypothesis\s+(one|two|three|four|five|six|\d+)",
            r"^demographic",
        ]
        return any(re.search(p, lowered) for p in patterns)

    for line in lines:
        stripped = line.strip()

        if is_table_start(stripped):
            if current_block:
                tables.append("\n".join(current_block).strip())
            current_block = pending_context + [stripped]
            pending_context = []
            continue

        if current_block:
            current_block.append(stripped)
        else:
            if is_context_heading(stripped):
                pending_context.append(stripped)
            elif pending_context and len(pending_context) < 8:
                pending_context.append(stripped)

    if current_block:
        tables.append("\n".join(current_block).strip())

    cleaned_tables: List[str] = []
    seen_numbers = set()

    for block in tables:
        first_table_line = extract_table_title_line(block)
        table_no = table_number_from_text(first_table_line)
        if table_no == 999:
            continue
        if table_no in seen_numbers:
            continue
        seen_numbers.add(table_no)
        cleaned_tables.append(block.strip())

    cleaned_tables.sort(key=table_number_from_text)
    return cleaned_tables


def detect_table_type(title: str, body: str) -> str:
    text = f"{title}\n{body}".lower()

    if any(k in text for k in ["gender", "age range", "qualification", "demographic", "respondent", "stakeholder category"]):
        return "Demographic"
    if any(k in text for k in ["pearson correlation", "ppmc", "relationship between", "correlation"]):
        return "PPMC"
    if any(k in text for k in ["t-test", "t value", "mean ratings of male and female"]):
        return "Independent t-test"
    if any(k in text for k in ["anova", "between groups", "within groups", "f-value", "analysis of variance"]):
        return "ANOVA"
    if any(k in text for k in ["model summary", "coefficient", "regression", "r square", "adjusted r square"]):
        return "Linear Regression"
    if any(k in text for k in ["mediation", "indirect effect", "sobel", "direct effect", "a path", "b path"]):
        return "Mediation"
    if any(k in text for k in ["mean", "std. deviation", "ranking", "decision", "weighted mean", "remarks"]):
        return "Descriptive"
    return "Other"


def infer_section_tag_from_context(text_block: str) -> str:
    lower = text_block.lower()

    rq_patterns = [
        r"research question\s+(one|two|three|four|five|six|\d+)\s*[:\.]?\s*([^\n]+)",
        r"rq\s+(one|two|three|four|five|six|\d+)\s*[:\.]?\s*([^\n]+)",
        r"question\s+(one|two|three|four|five|six|\d+)\s*[:\.]?\s*([^\n]+)",
    ]
    for pattern in rq_patterns:
        match = re.search(pattern, lower)
        if match:
            num = to_number(match.group(1))
            text = match.group(2).replace("?", "").strip()
            if num is not None and len(text) > 5:
                return f"RQ {num}: {text.title()}"

    h_patterns = [
        r"h0?(\d+)\s*[:\.]?\s*([^\n]+)",
        r"hypothesis\s+(one|two|three|four|five|six|\d+)\s*[:\.]?\s*([^\n]+)",
        r"research hypothesis\s+(one|two|three|four|five|six|\d+)\s*[:\.]?\s*([^\n]+)",
    ]
    for pattern in h_patterns:
        match = re.search(pattern, lower)
        if match:
            num = to_number(match.group(1))
            text = match.group(2).replace("?", "").strip()
            if num is not None and len(text) > 5:
                return f"H{str(num).zfill(2)}: {text.title()}"

    rq_match = re.search(r"research question\s+(one|two|three|four|five|six|\d+)", lower)
    if rq_match:
        num = to_number(rq_match.group(1))
        if num is not None:
            return f"RQ {num}"

    h_match = re.search(r"(?:h0?(\d+)|hypothesis\s+(one|two|three|four|five|six|\d+))", lower)
    if h_match:
        raw = h_match.group(1) or h_match.group(2)
        num = to_number(raw)
        if num is not None:
            return f"H{str(num).zfill(2)}"

    if "demographic" in lower or "gender" in lower or "age range" in lower:
        return "DEMOGRAPHIC"

    return "OTHER"


def parse_table_chunk(chunk: str) -> Dict[str, Any]:
    lines = [ln.strip() for ln in chunk.split("\n") if ln.strip()]

    table_line_index = None
    for idx, line in enumerate(lines):
        if extract_table_number(line) is not None:
            table_line_index = idx
            break

    if table_line_index is None:
        table_line_index = 0

    context_lines = lines[:table_line_index]
    title = lines[table_line_index] if lines else "Untitled Table"
    body = "\n".join(lines[table_line_index + 1:]) if len(lines) > table_line_index + 1 else ""
    context_text = "\n".join(context_lines)
    actual_table_number = table_number_from_text(title)

    return {
        "original_title": title,
        "body": body,
        "context_text": context_text,
        "table_type": detect_table_type(title, body),
        "section_tag": infer_section_tag_from_context(chunk),
        "chunk": chunk,
        "table_number": actual_table_number,
    }


def extract_source_line(chunk: str) -> str:
    match = re.search(r"(Source:\s*.*)", chunk, flags=re.I)
    return match.group(1).strip() if match else "Source: Researcher’s Output (Year) Using SPSS"


def build_apa_title(table: Dict[str, Any], index: int) -> str:
    original_title = table.get("original_title", "").strip()
    actual_number = table.get("table_number", table_number_from_text(original_title))
    if actual_number != 999:
        return original_title
    return f"Table {index}: {original_title}"


def normalize_rows(rows: List[List[str]]) -> List[List[str]]:
    if not rows:
        return []
    max_len = max(len(r) for r in rows)
    return [r + [""] * (max_len - len(r)) for r in rows]


def split_columns_from_line(line: str) -> List[str]:
    if "\t" in line:
        parts = [p.strip() for p in line.split("\t")]
        return [p for p in parts if p != ""]

    if "|" in line:
        parts = [p.strip() for p in line.split("|")]
        return [p for p in parts if p != ""]

    parts = [p.strip() for p in re.split(r"\s{2,}", line) if p.strip()]
    if len(parts) >= 2:
        return parts

    return [line.strip()] if line.strip() else []


def chunk_to_rows(chunk: str) -> List[List[str]]:
    lines = [ln.rstrip() for ln in chunk.split("\n") if ln.strip()]
    rows: List[List[str]] = []
    for line in lines:
        if re.match(
            r"^(Research Question|Question|RQ\s+\d+|Hypothesis|Research Hypothesis|H0?\d+|Source:)",
            line,
            flags=re.I,
        ):
            continue
        cols = split_columns_from_line(line)
        if cols:
            rows.append(cols)
    return normalize_rows(rows)


def detect_apa_note(table: Dict[str, Any]) -> str:
    chunk = table.get("chunk", "")
    lower = chunk.lower()

    if "correlation is significant" in lower:
        match = re.search(r"significant at the\s+([0-9.]+)\s+level", chunk, flags=re.I)
        if match:
            return f"Note. Correlation was significant at the {match.group(1)} level (2-tailed)."
        return "Note. Correlation was significant at the stated level (2-tailed)."

    if "p<0.05" in lower or "@p<0.05" in lower:
        return "Note. Decision was based on p < .05."

    if "p>0.05" in lower:
        return "Note. Decision was based on p > .05."

    return ""


def rows_to_markdown_table(rows: List[List[str]]) -> str:
    if not rows:
        return ""

    rows = normalize_rows(rows)
    if len(rows) == 1:
        header = [f"Column {i + 1}" for i in range(len(rows[0]))]
        body = rows
    else:
        header = rows[0]
        body = rows[1:]

    header_line = "| " + " | ".join(header) + " |"
    sep_line = "| " + " | ".join(["---"] * len(header)) + " |"
    body_lines = ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join([header_line, sep_line] + body_lines)


def build_apa_table_block(table: Dict[str, Any], index: int) -> str:
    rows = chunk_to_rows(table.get("chunk", ""))
    md_table = rows_to_markdown_table(rows)
    note = detect_apa_note(table)
    title = table.get("apa_title", f"Table {index}")
    source = table.get("source_line", "")

    block = [f"**{title}**", "", md_table or "Table data could not be formatted."]
    if note:
        block.extend(["", f"*{note}*"])
    if source:
        block.extend(["", f"*{source}*"])
    return "\n".join(block)


def sort_for_final_output(tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    def key_func(item: Dict[str, Any]):
        tag = str(item.get("section_tag", "OTHER"))
        if tag.startswith("RQ"):
            return (0, section_sort_number(tag), table_number_from_text(item.get("apa_title", "")))
        if tag.startswith("H"):
            return (1, section_sort_number(tag), table_number_from_text(item.get("apa_title", "")))
        if tag == "DEMOGRAPHIC":
            return (2, 0, table_number_from_text(item.get("apa_title", "")))
        return (3, 999, table_number_from_text(item.get("apa_title", "")))

    return sorted(tables, key=key_func)


# =========================================================
# PROMPTS
# =========================================================
def build_all_interpretations_prompt(tables: List[Dict[str, Any]]) -> str:
    blocks = []
    for i, table in enumerate(tables, start=1):
        table_type = table["table_type"].upper()

        format_examples = ""
        if table_type == "DEMOGRAPHIC":
            format_examples = """
EXAMPLE FORMAT FOR DEMOGRAPHIC:
Table 1 above described the categories of the respondents based on gender. It revealed that male respondents are 97 (48.85%) and female respondents are 103 (51.5%). This implies that there are more females than male in the selected schools.
Table 1 outlined the demographic distribution of the 200 respondents who participated in the study. It showed that 54.5% were males and 45.5% were females, indicating a slightly higher number of male respondents than female respondents. Regarding age range, the majority of the respondents (54.5%) were in the 41–50 years category, followed by those aged 31–40 years (33%). Respondents in the 18–30 and 50 and above categories were the least represented, with 7.5% and 5.0%, respectively. This suggests that most participants were middle-aged individuals, likely experienced in the education sector. In terms of educational qualifications, the majority of respondents held a Bachelor of Education (B.Ed.) degree (57%), followed by those with an NCE (27%). A smaller proportion possessed SSCE and M.Ed. qualifications (8% each). This indicated that most participants had a formal background in education. Concerning stakeholder categories, nearly half of the respondents (49.5%) were teachers, while 24% were parents, 7% were school administrators, and 19.5% fell into other categories. This demonstrated that the study included a diverse group of stakeholders involved in secondary education. Regarding years of involvement in secondary education, the largest group (45.5%) was involved for 5–10 years, followed by those with 11–15 years (32.5%) of experience. Respondents with less than 5 years accounted for 13%, while only 9% had more than 15 years of experience. This implied that most respondents accumulated a moderate to substantial level of experience in secondary education.
"""
        elif table_type == "DESCRIPTIVE":
            format_examples = """
EXAMPLE FORMAT FOR DESCRIPTIVE:
Table 2 disclosed how different stakeholders contributed to school development based on students' opinions. The highest-rated contribution was parents supporting the school financially, with a mean score of 4.78, ranked 1st, suggesting it was perceived as the most significant input. This was followed by school administrators providing strong leadership (4.70), ranked 2nd, and teachers participating in curriculum planning, which scored 4.68, ranked 3rd. Open communication between administrators and stakeholders (4.56) and teamwork between teachers and parents (4.54) were ranked 4th and 5th, respectively. Although still rated positively, community support for school buildings (4.45), parental involvement in learning (4.48), and parents helping in school decisions (4.49) were ranked slightly lower. The findings indicated that all stakeholders played supportive roles in school development, with financial support, leadership, and teacher involvement in curriculum planning standing out as the most impactful contributions.
Table 5 presented how public and private secondary schools in Ilorin Metropolis engaged with their communities in school development. The findings revealed that schools frequently organize meetings with parents to discuss development issues (M = 3.21, SD = 0.57, Accepted), indicating that parental involvement is a well-established practice. Similarly, schools often invite professionals from the community to mentor students (M = 2.68, SD = 0.99, Accepted), suggesting that mentorship programs are fairly common, though they could be strengthened. However, other forms of community involvement appear limited. Fundraising support from the community is rare (M = 2.14, SD = 0.93, Not accepted), meaning schools do not frequently receive financial assistance from community members. Likewise, collaboration with local businesses to support educational programs is infrequent (M = 2.33, SD = 0.97, Not accepted), highlighting a weak partnership between schools and businesses. Furthermore, schools rarely allow the community to use school facilities for events (M = 2.12, SD = 1.10, Not accepted), suggesting a restrictive policy on facility access. With an overall weighted mean of 2.5, the results indicate that while parental involvement and student mentorship are relatively common, financial support, business collaboration, and facility access remain underutilized opportunities for school development.
The above table revealed that 25 (20.8%) showed a low level of pupils' academic performance in basic science, 57 (47.5%) showed an average level of pupils' academic performance in basic science, and 38 (31.7%) showed a high level of pupils' academic performance in basic science. This revealed that the majority of the pupils' (47.5%) responses showed an average level of academic performance in basic science. This implies that primary school pupils in Ilorin Metropolis, Kwara State, have an average level of academic performance in basic science.
Table 1 presents the level of principals' professional skills in secondary schools in Southwest, Nigeria. The result depicts that, using a criterion mean score of 2.50 for the rating scale; all the items had mean scores above the cut-off point. This implies that the level of principals' professional skills in secondary schools in Southwest, Nigeria was High.
Table 3 present the level of academic achievement of secondary school students (cognitive domain). Table 4 shows that, out of 51,975 students who enrolled for Senior School Certificate Examination (SSCE) between 2015 to 2020; 28289 students representing 54.4% obtained 5 credits and above With English Language and Mathematics, 12335 (23.7%) had 5 credits and above Without English Language and Mathematics While 7396 (14.2%) and 3955 7.6%) students had 4 credits and less than 4 credits respectively. Therefore. Level of academic achievement of secondary school students in Southwest, Nigeria was moderate. The level of academic achievement of secondary school students was moderate on the basis of the number (28,289) and percentage (54.4%) of students who had five credits and above with English Language and Mathematics.
"""
        elif table_type == "PPMC":
            format_examples = """
EXAMPLE FORMAT FOR PPMC:
Table 7 presented the result of the analysis examining the relationship between principals' management practices and teacher retention rate in secondary schools in Ilorin South. The result showed a moderate positive correlation (r = 0.38) between management practices and teacher retention, with a p-value of 0.01, which is less than 0.05. This indicates that the result is statistically significant. Therefore, there is sufficient evidence to reject the null hypothesis and conclude that there is a meaningful relationship between the management practices of school principals and teacher retention rates. In simpler terms, schools where principals exhibit strong management practices are more likely to retain their teachers.
"""
        elif table_type == "INDEPENDENT T-TEST":
            format_examples = """
EXAMPLE FORMAT FOR INDEPENDENT T-TEST:
Table 5 showed that a t-value of 1.51 was obtained with a p-value of 0.13, which was greater than the 0.05 level of significance. Consequently, the null hypothesis was retained. This indicated that there was no statistically significant difference in the mean ratings of male and female respondents on the relevance of SIWES to the development of students' academic careers. The null hypothesis was therefore retained.
"""
        elif table_type == "ANOVA":
            format_examples = """
EXAMPLE FORMAT FOR ANOVA:
As shown in Table 4, the F-value of 1.283 with a p-value of 0.23 computed at 0.05 alpha level. Since the p-value of 0.23 obtained is greater than 0.05 level of significance, the null hypothesis two is retained.
"""
        elif table_type == "LINEAR REGRESSION":
            format_examples = """
EXAMPLE FORMAT FOR LINEAR REGRESSION:
The model summary tests the null hypothesis that there is no significant impact of age on teacher attrition rates in private secondary schools. The results indicate that the correlation coefficient is nearly zero, signifying a negligible linear relationship between age and teacher attrition rates. The coefficient of determination (R Square) is also zero, suggesting that age accounts for virtually none of the variance in teacher attrition rates. The negative adjusted R Square value further indicates that the model does not fit the data well. Additionally, the standard error of the estimate is relatively large at 0.74965, reflecting a substantial average distance between the observed values and the regression line.
The change statistics reinforce these findings, showing no change in R Square and an F value of zero, indicating no improvement in the model's fit when age is included as a predictor. The significance value for the F change is 0.987, far above the conventional threshold for statistical significance, meaning that the impact of age on teacher attrition rates is not statistically significant.
The model summary clearly shows that age does not significantly impact teacher attrition rates in private secondary schools. This is supported by the low correlation and determination coefficients, the negative adjusted R Square, and the high significance value, all of which validate the null hypothesis.
The ANOVA table evaluates the hypothesis that there is no significant impact of age on teacher attrition rates in private secondary schools. The table includes the sum of squares, degrees of freedom (df), mean squares, F-statistic, and significance level (Sig.).
In the regression row, the sum of squares is 0.000, indicating that the variability explained by age is virtually nonexistent. With 1 degree of freedom, the mean square for the regression is also 0.000, further suggesting no variation due to age. The F-statistic is 0.000, reinforcing that age does not explain any variability in teacher attrition rates.
The residual row represents the unexplained variability. The sum of squares for the residual is 202.875, with 361 degrees of freedom, resulting in a mean square of 0.562. This high residual sum of squares compared to the regression sum of squares indicates that most of the variability in teacher attrition rates is unexplained by age.
The significance level (Sig.) is 0.987, which is much higher than the conventional threshold of 0.05. This means that there is no statistically significant relationship between age and teacher attrition rates.
The ANOVA results confirm that age has no significant impact on teacher attrition rates in private secondary schools. The regression sum of squares is negligible, the F-statistic is zero, and the high significance value supports the null hypothesis.
The coefficients table evaluates the impact of age on teacher attrition rates in private secondary schools. The table provides values for unstandardized coefficients, standardized coefficients, t-statistic, and significance level (Sig.) for the model.
The constant (intercept) has an unstandardized coefficient (B) of 4.151 with a standard error of 0.169. This value represents the expected teacher attrition rate when the age variable is zero. The t-statistic for the constant is 24.626, with a significance level (Sig.) of 0.000, indicating that the constant is highly significant.
For the age variable, the unstandardized coefficient (B) is 0.001, with a standard error of 0.055. This coefficient indicates that for each one-year increase in age, the teacher attrition rate changes by 0.001 units. The standardized coefficient (Beta) is also 0.001, reflecting a negligible effect of age on teacher attrition rates when accounting for the scale of measurement.
The t-statistic for the age variable is 0.016, with a significance level (Sig.) of 0.987. This high p-value indicates that the age variable is not statistically significant in predicting teacher attrition rates.
The coefficients table shows that the age of teachers has no significant impact on their attrition rates in private secondary schools. The unstandardized and standardized coefficients for age are close to zero, and the high p-value supports the null hypothesis that age does not affect teacher attrition.
"""
        elif table_type == "MEDIATION":
            format_examples = """
EXAMPLE FORMAT FOR MEDIATION:
The mediation analysis revealed that Health Impact significantly predicted Social Support. Finally, the Sobel test did not confirm mediation (z = 0.80, p > .05), indicating that Social Support did not significantly mediate the relationship between Health Impact and Social Isolation. These results suggested that the relationship between Health Impact and Social Isolation was not mediated by Social Support.
"""

        blocks.append(
            f"""
###TABLE_{i}_START###
TABLE TYPE: {table_type}
APA TITLE: {table['apa_title']}
SECTION TAG: {table['section_tag']}
SOURCE LINE: {table['source_line']}
CONTEXT TEXT:
{table.get('context_text', '')}
RAW TABLE BODY:
{table['body']}
{format_examples}
###TABLE_{i}_END###
"""
        )

    return f"""
STRICT EXECUTION MODE – FOLLOW MY INTERPRETATION GUIDE EXACTLY.

For EACH table, provide interpretation using the specified format for its table type.

CRITICAL RULES:
- Use exact format examples provided for each table type
- Write in past tense throughout
- For DEMOGRAPHIC tables: describe respondent characteristics and distributions
- For DESCRIPTIVE tables: include statistical values (means, SD, rankings) and explain what they mean
- For INFERENTIAL tables (PPMC, t-test, ANOVA, Regression, Mediation): include statistical results (p-values, correlations, F-values) but NOT mean/SD unless they are part of the test results
- Use varied phrases like "revealed that", "showed that", "indicated that", "implied that", "presented that", "demonstrated that", "displayed that", "found that"
- Keep each interpretation to one comprehensive paragraph
- Do NOT generate APA tables - only provide interpretations
- Wrap each interpretation in these markers:
  ###ANALYSIS_1### ... ###END_ANALYSIS_1###
  ###ANALYSIS_2### ... ###END_ANALYSIS_2###

TABLES TO INTERPRET:
{''.join(blocks)}
"""


def parse_bulk_interpretations(result: str, tables: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    output = []
    for i, table in enumerate(tables, start=1):
        pattern = rf"###ANALYSIS_{i}###(.*?)###END_ANALYSIS_{i}###"
        match = re.search(pattern, result, flags=re.S)
        interpretation = match.group(1).strip() if match else "Analysis could not be parsed."
        item = dict(table)
        item["interpretation"] = interpretation
        output.append(item)
    return output


def build_summary_prompt(all_tables: List[Dict[str, Any]]) -> str:
    total_tables = len(all_tables)

    joined = "\n\n".join(
        [
            f"TABLE: {t['apa_title']}\nINTERPRETATION: {t['interpretation']}"
            for t in all_tables
        ]
    )

    return f"""
Generate a summary of findings based only on the interpreted tables below.

Rules:
- Start with exactly: The summary of findings based on the tables is presented below:
- Use numbered bullet points (1., 2., 3., etc.)
- Base the summary on the tables directly, using Table 1, Table 2, Table 3, Table 4, Table 5, etc.
- Keep language simple and direct
- Each point should be a complete sentence that captures a key finding from the corresponding table
- Do not add statistical details in the summary - focus on the main conclusions
- Use all {total_tables} tables
- Do not group the summary under research questions or hypotheses
- Do not separate demographic findings from other tables

INTERPRETED FINDINGS:
{joined}
"""


def build_discussion_prompt(non_demo_tables: List[Dict[str, Any]], empirical_review: str) -> str:
    sorted_tables = sorted(non_demo_tables, key=lambda x: table_number_from_text(x.get("apa_title", "")))
    findings_text = ""
    for i, table in enumerate(sorted_tables, start=1):
        table_number_match = re.search(r"Table\s+(\d+)", table["apa_title"], re.IGNORECASE)
        actual_table_number = table_number_match.group(1) if table_number_match else str(i)
        findings_text += f"Table {actual_table_number} ({table['apa_title']}): {table['interpretation']}\n\n"

    return f"""
You are an expert academic writing assistant.
Generate DISCUSSION OF FINDINGS using the exact format provided below.

Rules:
1. Use "Findings from research question one revealed that", "Findings from Research Question Two indicated that", "Findings from Research Question Three showed that", etc. for research questions
2. Use "Findings from Hypothesis One showed that", "Findings from Hypothesis Two revealed that", etc. for hypotheses
3. For each table, write exactly one comprehensive paragraph
4. Compare findings with previous studies from the empirical review only
5. Use exactly one citation per paragraph from the empirical review only
6. Add a REFERENCES section at the end
7. Do not use chapter one, chapter three, methodology, or other non-empirical sections
8. Follow this exact format for each paragraph:

"Findings from research question one revealed that [main finding]. [Additional details]. This aligns with the work of [Author(s)] ([Year]), who found that [comparison finding]. While both studies agree that [common point], the present study uniquely highlights [unique contribution] in the Nigerian context."

"Findings from Research Question Two indicated that [main finding]. [Additional details]. These results are supported by [Author(s)] ([Year]), who found [comparison finding]. While their study emphasized [their focus], the present study narrows the lens to [current focus], offering a context-specific contribution to Nigerian secondary education research."

"Findings from Research Question Three showed that [main finding]. [Additional details]. This corresponds with [Author(s)] ([Year]), who also found [comparison finding]. Both studies affirm that [common finding]; however, while [Author(s)] focused on [their focus], this study expands generalizability by focusing on [current focus], capturing broader patterns in [broader context]."

"Findings from Research Question Four revealed that [main finding]. [Additional details]. This somewhat aligns with [Author(s)] ([Year]), whose work suggested [comparison finding]. Though [Author(s)] focused more on [their focus], both studies acknowledge that [common issue]."

"Findings from Research Question Five showed that [main finding]. [Additional details]. This is consistent with [Author(s)] ([Year]), who found that [comparison finding]. However, the current study innovates by [unique contribution], suggesting that [unique insight]."

"Findings from Hypothesis One showed [main finding]. This contrasts with [Author(s)] ([Year]), who found [comparison finding]. The discrepancy may be explained by [explanation]."

"Findings from Hypothesis Two revealed [main finding]. This is similar to [Author(s)] ([Year]), who found [comparison finding]. Both studies reveal that [common point]. Yet, this study uniquely emphasizes [unique emphasis]."

"Findings from Hypothesis Three demonstrated [main finding]. This finding echoes [Author(s)] ([Year]), whose [analysis type] showed [comparison finding]. Both findings highlight that [common conclusion]. However, while [Author(s)] emphasized [their emphasis], the present study strengthens the argument by [current contribution]."

FINDINGS TO DISCUSS:
{findings_text}

EMPIRICAL REVIEW ONLY:
{empirical_review[:15000]}
"""


def build_final_outputs_prompt(non_demo_tables: List[Dict[str, Any]], study_context_text: str = "") -> str:
    sorted_tables = sort_for_final_output(non_demo_tables)
    findings_text = ""
    for i, table in enumerate(sorted_tables, start=1):
        findings_text += f"Table {i} ({table['apa_title']}): {table['interpretation']}\n\n"

    return f"""
Generate only these sections in this exact order:

1. CONCLUSION
2. IMPLICATION OF THE STUDY
3. RECOMMENDATIONS

Rules:
- The conclusion must be exactly 200 words in one paragraph
- Base everything primarily on the findings below
- Use the study context only to infer the setting, variables, likely respondents, and scope
- Do not invent data values
- Give a minimum of 7 implications of the study
- Give a minimum of 7 recommendations
- Make the implication of the study practical and specific to the study findings
- Use clear headings for each section
- Follow academic writing conventions
- Keep implications and recommendations simple and concise

FINDINGS TO USE:
{findings_text}

STUDY CONTEXT:
{study_context_text[:12000]}
"""


def build_comprehensive_report_prompt(non_demo_tables: List[Dict[str, Any]], empirical_review: str, study_context_text: str = "") -> str:
    sorted_tables = sort_for_final_output(non_demo_tables)
    findings_text = ""
    for i, table in enumerate(sorted_tables, start=1):
        findings_text += f"Table {i} ({table['apa_title']}): {table['interpretation']}\n\n"

    return f"""
Generate a comprehensive academic report with the following sections in this exact order:

1. DISCUSSION OF FINDINGS
2. CONCLUSION
3. IMPLICATION OF THE STUDY
4. RECOMMENDATIONS
5. LIMITATION OF THE STUDY
6. SUGGESTIONS FOR FUTURE RESEARCH

Rules for DISCUSSION OF FINDINGS:
- Write exactly one paragraph for each interpreted non-demographic table
- Use exactly one citation per paragraph from the empirical review only
- If you cannot find a citation in the empirical review that truly aligns with the finding, do NOT invent one
- Instead write clearly: "No directly aligned citation was found in the empirical review for this finding."
- Add a REFERENCES section at the end using only studies actually cited in the discussion
- Do not fabricate authors, years, journal names, or assertions

Rules for CONCLUSION:
- Exactly 200 words in one paragraph
- Summarize the major findings only

Rules for IMPLICATION OF THE STUDY:
- Numbered points
- Practical and concise
- Based on the findings and study context
- Minimum of 7 implications

Rules for RECOMMENDATIONS:
- Numbered points
- Minimum of 7 recommendations
- Based on the findings

Rules for LIMITATION OF THE STUDY:
- Numbered points
- Minimum of 7 limitations
- Use the study context and the research questions where appropriate
- Be realistic and concise

Rules for SUGGESTIONS FOR FUTURE RESEARCH:
- Numbered points
- Minimum of 7 suggestions
- Use the research questions and study context where appropriate
- Be specific and concise

FINDINGS TO USE:
{findings_text}

EMPIRICAL REVIEW ONLY:
{empirical_review[:15000]}

STUDY CONTEXT:
{study_context_text[:12000]}
"""


# =========================================================
# EMPIRICAL REVIEW EXTRACTION
# =========================================================
def extract_empirical_review_section(full_text: str) -> str:
    text = clean_text(full_text)
    if not text:
        return ""

    start_patterns = [
        r"(?im)^\s*empirical review of\b.*$",
        r"(?im)^\s*empirical review\b.*$",
        r"(?im)^\s*review of empirical studies\b.*$",
        r"(?im)^\s*empirical studies\b.*$",
    ]

    end_patterns = [
        r"(?im)^\s*summary of reviewed literature\b.*$",
        r"(?im)^\s*appraisal of reviewed literature\b.*$",
        r"(?im)^\s*chapter\s+three\b.*$",
        r"(?im)^\s*research methodology\b.*$",
        r"(?im)^\s*methodology\b.*$",
    ]

    start_match = None
    for pattern in start_patterns:
        match = re.search(pattern, text)
        if match:
            start_match = match
            break

    if start_match:
        sliced = text[start_match.start():]
        end_match = None
        for pattern in end_patterns:
            match = re.search(pattern, sliced)
            if match and match.start() > 0:
                if end_match is None or match.start() < end_match.start():
                    end_match = match
        if end_match:
            return sliced[:end_match.start()].strip()
        return sliced.strip()

    return text.strip()


def empirical_review_has_citation_markers(text: str) -> bool:
    if not text:
        return False

    patterns = [
        r"\([A-Z][A-Za-z]+,\s*\d{4}\)",
        r"\([A-Z][A-Za-z]+\s+and\s+[A-Z][A-Za-z]+,\s*\d{4}\)",
        r"\([A-Z][A-Za-z]+\s+et al\.\s*\d{4}\)",
        r"\b[A-Z][A-Za-z]+\s*\(\d{4}\)",
        r"\b[A-Z][A-Za-z]+\s+and\s+[A-Z][A-Za-z]+\s*\(\d{4}\)",
        r"\b[A-Z][A-Za-z]+\s+et al\.\s*\(\d{4}\)",
    ]
    return any(re.search(p, text) for p in patterns)


# =========================================================
# REPORT EXPORT
# =========================================================
def export_full_report(
    all_tables: List[Dict[str, Any]],
    summary_text: str,
    discussion_text: str,
    final_output_text: str,
    comprehensive_text: str = "",
) -> str:
    lines = [f"# {APP_NAME}\n"]

    lines.append("## TABLE INTERPRETATIONS\n")
    for i, table in enumerate(all_tables, start=1):
        lines.append(f"**{table['apa_title']}**")
        lines.append("")
        lines.append(table.get("interpretation", ""))
        lines.append("")

    if summary_text.strip():
        lines.append("## SUMMARY OF FINDINGS\n")
        lines.append(summary_text)
        lines.append("")

    if discussion_text.strip():
        lines.append("## DISCUSSION OF FINDINGS\n")
        lines.append(discussion_text)
        lines.append("")

    if final_output_text.strip():
        lines.append("## CONCLUSION, IMPLICATION, RECOMMENDATIONS\n")
        lines.append(final_output_text)
        lines.append("")

    if comprehensive_text.strip():
        lines.append("## LIMITATIONS AND FUTURE RESEARCH\n")
        lines.append(comprehensive_text)

    return "\n".join(lines)


def create_text_download_docx(content: str) -> BytesIO:
    doc = Document()

    title = doc.add_heading(APP_NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = content.split("\n")
    for line in lines:
        if line.strip():
            if line.startswith("###"):
                heading = line.replace("###", "").strip()
                doc.add_heading(heading, level=1)
            else:
                doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


def build_complete_report_content() -> str:
    combined_content = ""

    combined_content += "### TABLE INTERPRETATIONS\n\n"
    for item in st.session_state.interpreted_tables:
        combined_content += f"**{item['apa_title']}**\n\n"
        combined_content += item["interpretation"] + "\n\n"

    if st.session_state.summary_of_findings:
        combined_content += "### SUMMARY OF FINDINGS\n\n"
        combined_content += st.session_state.summary_of_findings + "\n\n"

    if st.session_state.get("discussion_of_findings", ""):
        combined_content += "### DISCUSSION OF FINDINGS\n\n"
        combined_content += st.session_state.discussion_of_findings + "\n\n"

    if st.session_state.get("final_outputs", ""):
        combined_content += "### CONCLUSION, IMPLICATION, RECOMMENDATIONS\n\n"
        combined_content += st.session_state.final_outputs + "\n\n"

    if st.session_state.limitations_future_research:
        combined_content += "### LIMITATIONS AND FUTURE RESEARCH\n\n"
        combined_content += st.session_state.limitations_future_research.strip() + "\n\n"

    if st.session_state.comprehensive_report:
        combined_content = st.session_state.comprehensive_report.strip()

    return combined_content.strip()


def build_full_reports_content() -> str:
    combined_content = ""

    if st.session_state.interpreted_tables:
        combined_content += "### ALL INTERPRETATIONS\n\n"
        for item in st.session_state.interpreted_tables:
            combined_content += f"**{item['apa_title']}**\n\n"
            combined_content += item.get("interpretation", "").strip() + "\n\n"

    if st.session_state.summary_of_findings:
        combined_content += "### SUMMARY OF FINDINGS\n\n"
        combined_content += st.session_state.summary_of_findings.strip() + "\n\n"

    if st.session_state.discussion_of_findings:
        combined_content += "### DISCUSSION OF FINDINGS\n\n"
        combined_content += st.session_state.discussion_of_findings.strip() + "\n\n"

    if st.session_state.final_outputs:
        combined_content += "### CONCLUSION, IMPLICATION OF THE STUDY, AND RECOMMENDATIONS\n\n"
        combined_content += st.session_state.final_outputs.strip() + "\n\n"

    if st.session_state.limitations_future_research:
        combined_content += "### LIMITATION OF THE STUDY AND SUGGESTIONS FOR FUTURE RESEARCH\n\n"
        combined_content += st.session_state.limitations_future_research.strip() + "\n\n"

    return combined_content.strip()


def create_text_download_pdf(content: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        alignment=1,
        spaceAfter=30,
    )

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=12,
    )

    story = []
    story.append(Paragraph(APP_NAME, title_style))
    story.append(Spacer(1, 20))

    lines = content.split("\n")
    for line in lines:
        if line.strip():
            if line.startswith("###"):
                heading = line.replace("###", "").strip()
                story.append(Paragraph(heading, heading_style))
                story.append(Spacer(1, 12))
            else:
                story.append(Paragraph(line.strip(), styles["Normal"]))
                story.append(Spacer(1, 6))
        else:
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer


def create_docx_report(
    all_tables: List[Dict[str, Any]],
    summary_text: str,
    discussion_text: str,
    final_output_text: str,
    comprehensive_text: str = "",
) -> BytesIO:
    doc = Document()

    title = doc.add_heading(APP_NAME, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary_intro = doc.add_paragraph()
    summary_intro.add_run(f"Summary of Findings Based on {len(all_tables)} Interpreted Tables").bold = True
    summary_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if comprehensive_text:
        add_formatted_text_to_docx(doc, comprehensive_text)
    else:
        if summary_text.strip():
            doc.add_heading("SUMMARY OF FINDINGS", level=1)
            add_formatted_text_to_docx(doc, summary_text)

        if discussion_text.strip():
            doc.add_heading("DISCUSSION OF FINDINGS", level=1)
            add_formatted_text_to_docx(doc, discussion_text)

        if final_output_text.strip():
            doc.add_heading("CONCLUSION, IMPLICATION, RECOMMENDATIONS, LIMITATION AND FUTURE RESEARCH", level=1)
            add_formatted_text_to_docx(doc, final_output_text)

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


def add_formatted_text_to_docx(doc, text: str):
    lines = text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.isupper() and len(line) < 50 and not line.endswith("."):
            doc.add_heading(line, level=2)
        elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        elif line.startswith(("CONCLUSION", "IMPLICATION", "RECOMMENDATIONS", "LIMITATION", "SUGGESTIONS")):
            doc.add_heading(line, level=1)
        elif line.startswith("REFERENCES"):
            doc.add_heading(line, level=1)
        else:
            p = doc.add_paragraph()
            p.add_run(line)


def create_pdf_report(
    all_tables: List[Dict[str, Any]],
    summary_text: str,
    discussion_text: str,
    final_output_text: str,
    comprehensive_text: str = "",
) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        alignment=1,
        spaceAfter=30,
    )

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=12,
    )

    story = []
    story.append(Paragraph(APP_NAME, title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Summary of Findings Based on {len(all_tables)} Interpreted Tables", heading_style))
    story.append(Spacer(1, 20))

    content_text = comprehensive_text if comprehensive_text else f"\n\n{summary_text}\n\n{discussion_text}\n\n{final_output_text}"

    paragraphs = content_text.split("\n\n")
    for para in paragraphs:
        if para.strip():
            if para.strip().isupper() and len(para.strip()) < 50:
                story.append(Paragraph(para.strip(), heading_style))
            elif para.strip().startswith(("CONCLUSION", "IMPLICATION", "RECOMMENDATIONS", "LIMITATION", "SUGGESTIONS", "REFERENCES")):
                story.append(Paragraph(para.strip(), heading_style))
            else:
                lines = para.split("\n")
                for line in lines:
                    if line.strip():
                        if line.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                            formatted_line = f"<b>{line.strip()}</b>"
                            story.append(Paragraph(formatted_line, styles["Normal"]))
                        else:
                            story.append(Paragraph(line.strip(), styles["Normal"]))
                    else:
                        story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer


# =========================================================
# INIT
# =========================================================
init_db()


# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.markdown("## ⚙️ Control Room")

api_key = get_api_key()
manual_key = st.sidebar.text_input("Gemini API Key", type="password")
if manual_key.strip():
    api_key = manual_key.strip()

model_name = st.sidebar.selectbox(
    "Gemini Model",
    [DEFAULT_MODEL, "gemini-2.5-pro", "gemini-2.5-flash-lite"],
    index=0,
)

if st.sidebar.button("Clear Current Session"):
    for key, val in defaults.items():
        st.session_state[key] = val
    st.rerun()


# =========================================================
# HEADER
# =========================================================
st.markdown("<div class='shell'>", unsafe_allow_html=True)

logo_col, title_col = st.columns([1, 3])
with logo_col:
    if os.path.exists(LOGO_PATH):
        try:
            img = Image.open(LOGO_PATH)
            st.image(img, use_container_width=True)
        except Exception:
            st.warning("Logo file could not be displayed.")
    else:
        st.info("Put 'ABT LOGO.jpg' in the same folder as this app.")

with title_col:
    st.markdown(
        f"""
    <div class="hero">
        <div class="hero-title">{APP_NAME}</div>
        <div class="hero-sub">
            Upload all your tables at once, generate APA-style tables, interpret everything at once,
            then produce summary of findings, discussion of findings, conclusion, implication,
            recommendations, limitation, and future research in one workflow.
        </div>
        <div class="badge-row">
            <div class="badge-pill">DOCX + PDF Upload</div>
            <div class="badge-pill">Interpret All At Once</div>
            <div class="badge-pill">APA-Style Tables</div>
            <div class="badge-pill">Summary First</div>
            <div class="badge-pill">Full Report Option</div>
            <div class="badge-pill">Saved Histories</div>
        </div>
    </div>
    """,
        unsafe_allow_html=True,
    )

c1, c2, c3 = st.columns(3)
with c1:
    st.markdown(
        f"""
    <div class="stat-card">
        <div class="stat-label">Detected Tables</div>
        <div class="stat-value">{len(st.session_state.parsed_tables)}</div>
    </div>
    """,
        unsafe_allow_html=True,
    )
with c2:
    st.markdown(
        f"""
    <div class="stat-card">
        <div class="stat-label">Interpreted Tables</div>
        <div class="stat-value">{len(st.session_state.interpreted_tables)}</div>
    </div>
    """,
        unsafe_allow_html=True,
    )
with c3:
    st.markdown(
        f"""
    <div class="stat-card">
        <div class="stat-label">Saved Histories</div>
        <div class="stat-value">{len(load_history(500))}</div>
    </div>
    """,
        unsafe_allow_html=True,
    )

if not api_key:
    st.warning("Add your GEMINI_API_KEY in Streamlit secrets or paste it in the sidebar.")

tab1, tab2, tab3, tab4 = st.tabs([
    "📥 Load Data",
    "📊 Tables & Reports",
    "🧠 Empirical Review",
    "🗂 Saved Histories",
])


# =========================================================
# TAB 1 - LOAD DATA
# =========================================================
with tab1:
    st.markdown("<div class='section-card'>", unsafe_allow_html=True)
    st.subheader("Load All Analysis Tables")

    upload_file = st.file_uploader(
        "Upload DOOCX or PDF containing your analysis tables (drag and drop file here)",
        type=["docx", "pdf"],
    )

    pasted_text = st.text_area(
        "paste all your tables here",
        height=320,
        placeholder="Paste all your tables here...",
    )

    if st.button("Load and Detect All Tables"):
        combined = ""

        if upload_file is not None:
            extracted = extract_text_from_upload(upload_file)
            if extracted.startswith("ERROR:"):
                st.error(extracted)
            else:
                combined += extracted + "\n\n"

        if pasted_text.strip():
            combined += pasted_text.strip()

        combined = clean_text(combined)

        if not combined:
            st.warning("Upload a file or paste your tables first.")
        else:
            chunks = split_tables_from_text(combined)
            parsed = [parse_table_chunk(c) for c in chunks]
            parsed = sorted(parsed, key=lambda x: x.get("table_number", 999))

            for i, table in enumerate(parsed, start=1):
                table["apa_title"] = build_apa_title(table, i)
                table["source_line"] = extract_source_line(table["chunk"])

            st.session_state.raw_analysis_text = combined
            st.session_state.parsed_tables = parsed
            st.session_state.interpreted_tables = []
            st.session_state.summary_of_findings = ""
            st.session_state.discussion_of_findings = ""
            st.session_state.final_outputs = ""
            st.session_state.full_report = ""
            st.session_state.comprehensive_report = ""
            st.session_state.limitations_future_research = ""

            detected_numbers = [
                t.get("table_number")
                for t in parsed
                if isinstance(t.get("table_number"), int) and t.get("table_number") != 999
            ]

            if detected_numbers:
                st.success(
                    f"{len(parsed)} tables were detected using the table labels/numbers found in your document "
                    f"(Table {min(detected_numbers)} to Table {max(detected_numbers)})."
                )
            else:
                st.success(f"{len(parsed)} tables were detected.")

            if api_key and parsed:
                with st.spinner("Auto-interpreting all tables..."):
                    prompt = build_all_interpretations_prompt(parsed)
                    raw_interpretations = gemini_text(api_key, prompt, model_name)
                    if raw_interpretations.startswith("ERROR:"):
                        st.error(raw_interpretations)
                    else:
                        interpreted = parse_bulk_interpretations(raw_interpretations, parsed)
                        st.session_state.interpreted_tables = interpreted
                        st.success("All tables interpreted successfully!")

    if st.session_state.raw_analysis_text:
        with st.expander("Preview Loaded Content"):
            st.text_area("Loaded Text", st.session_state.raw_analysis_text[:20000], height=320)

    st.markdown("</div>", unsafe_allow_html=True)


# =========================================================
# TAB 2 - TABLES AND REPORTS
# =========================================================
with tab2:
    st.markdown("<div class='section-card'>", unsafe_allow_html=True)
    st.subheader("Generate All Interpretations and Final Outputs")

    if not st.session_state.parsed_tables:
        st.info("Load your tables first in the Load Data tab.")
    else:
        st.markdown(
            """
        <div class='info-box'>
        This section displays interpreted tables, summary of findings, discussion of findings,
        conclusion, implication of the study, recommendations, limitation of the study,
        future research suggestions, and a downloadable report.
        </div>
        """,
            unsafe_allow_html=True,
        )

        if st.session_state.interpreted_tables:
            st.markdown("### TABLE INTERPRETATIONS")
            for item in st.session_state.interpreted_tables:
                with st.expander(item["apa_title"]):
                    st.markdown(item["interpretation"])

            if st.button("📋 Generate Summary of Findings", use_container_width=True):
                if not api_key:
                    st.error("Gemini API key is missing.")
                else:
                    with st.spinner("Generating summary of findings..."):
                        prompt = build_summary_prompt(st.session_state.interpreted_tables)
                        summary_text = gemini_text(api_key, prompt, model_name)
                        if summary_text.startswith("ERROR:"):
                            st.error(summary_text)
                        else:
                            st.session_state.summary_of_findings = summary_text
                            st.success("Summary of findings generated successfully!")

        if st.session_state.summary_of_findings:
            st.markdown("### SUMMARY OF FINDINGS")
            st.markdown(st.session_state.summary_of_findings)

    st.markdown("</div>", unsafe_allow_html=True)


# =========================================================
# TAB 3 - EMPIRICAL REVIEW
# =========================================================
with tab3:
    st.markdown("<div class='section-card'>", unsafe_allow_html=True)
    st.subheader("Load Empirical Review / Chapter 1–3")

    st.markdown(
        """
    <div class='info-box'>
    If a document is uploaded here, the app will read the full file but use only the Empirical Review section
    to generate the Discussion of Findings. The pasted text box below accepts only empirical review text.
    The study context will still be kept for generating conclusion, implication of the study, recommendations,
    limitation of the study, and suggestions for future research.
    </div>
    """,
        unsafe_allow_html=True,
    )

    chapter_file = st.file_uploader(
        "Upload Chapter 1–3 or document containing the empirical review (DOCX or PDF)",
        type=["docx", "pdf"],
        key="empirical_upload",
    )

    empirical_text = st.text_area(
        "Paste only empirical review text here",
        height=280,
        placeholder="Paste only the empirical review here.",
    )

    if st.button("Load Empirical Review"):
        combined = ""

        if chapter_file is not None:
            extracted = extract_text_from_upload(chapter_file)
            if extracted.startswith("ERROR:"):
                st.error(extracted)
            else:
                combined += extracted + "\n\n"

        if empirical_text.strip():
            combined += empirical_text.strip()

        combined = clean_text(combined)

        if not combined:
            st.warning("Upload a document or paste the empirical review first.")
        else:
            empirical_only = extract_empirical_review_section(combined)
            st.session_state.study_context_text = combined
            st.session_state.empirical_review_text = empirical_only
            st.success("Document loaded successfully. Only the empirical review will be used for discussion of findings.")

            if not empirical_review_has_citation_markers(empirical_only):
                st.warning("The extracted empirical review does not appear to contain clear citation markers. Discussion may state that no directly aligned citation was found where necessary.")

    if st.session_state.empirical_review_text:
        with st.expander("Preview Extracted Empirical Review"):
            st.text_area(
                "Empirical Review Preview",
                st.session_state.empirical_review_text[:20000],
                height=320,
            )

        with st.expander("Preview Full Study Context"):
            st.text_area(
                "Full Uploaded Document Preview",
                st.session_state.study_context_text[:20000],
                height=320,
            )

        if st.session_state.interpreted_tables:
            st.markdown("### GENERATE COMPLETE REPORT SECTIONS")

            col1, col2, col3, col4 = st.columns(4)

            with col1:
                if st.button("📚 Discussion of Findings", use_container_width=True):
                    if not api_key:
                        st.error("Gemini API key is missing.")
                    else:
                        non_demo = [
                            t for t in st.session_state.interpreted_tables
                            if t["section_tag"] != "DEMOGRAPHIC"
                        ]
                        if non_demo:
                            with st.spinner("Generating discussion of findings."):
                                prompt = build_discussion_prompt(non_demo, st.session_state.empirical_review_text)
                                text = gemini_text(api_key, prompt, model_name)
                                if text.startswith("ERROR:"):
                                    st.error(text)
                                else:
                                    st.session_state.discussion_of_findings = text
                                    st.success("Discussion of findings generated successfully!")
                        else:
                            st.error("No non-demographic tables found for discussion generation.")

            with col2:
                if st.button("📝 Conclusion, Implications of Study & Recommendations", use_container_width=True):
                    if not api_key:
                        st.error("Gemini API key is missing.")
                    else:
                        non_demo = [
                            t for t in st.session_state.interpreted_tables
                            if t["section_tag"] != "DEMOGRAPHIC"
                        ]
                        if non_demo:
                            with st.spinner("Generating conclusion, implication, and recommendations."):
                                prompt = build_final_outputs_prompt(non_demo, st.session_state.study_context_text)
                                text = gemini_text(api_key, prompt, model_name)
                                if text.startswith("ERROR:"):
                                    st.error(text)
                                else:
                                    st.session_state.final_outputs = text
                                    st.success("Conclusion, implication of the study, and recommendations generated successfully!")
                        else:
                            st.error("No non-demographic tables found for final outputs generation.")

            with col3:
                if st.button("🔍 Limitations of the Study and Suggestions for Future Research", use_container_width=True):
                    if not api_key:
                        st.error("Gemini API key is missing.")
                    else:
                        non_demo = [
                            t for t in st.session_state.interpreted_tables
                            if t["section_tag"] != "DEMOGRAPHIC"
                        ]
                        if non_demo:
                            with st.spinner("Generating limitation of the study and suggestions for future research."):
                                findings_text = ""
                                sorted_tables = sort_for_final_output(non_demo)
                                for i, table in enumerate(sorted_tables, start=1):
                                    findings_text += f"Table {i} ({table['apa_title']}): {table['interpretation']}\n\n"

                                prompt = f"""
Generate only these sections in this exact order:

1. LIMITATION OF THE STUDY
2. SUGGESTIONS FOR FUTURE RESEARCH

Rules:
- Base everything on the findings below
- Use the study context and the research questions where appropriate
- Provide a minimum of 7 substantial limitations
- Provide a minimum of 7 specific future research suggestions
- Use clear headings
- Be practical, realistic, and concise

FINDINGS TO USE:
{findings_text}

STUDY CONTEXT:
{st.session_state.study_context_text[:12000]}
"""
                                text = gemini_text(api_key, prompt, model_name)
                                if text.startswith("ERROR:"):
                                    st.error(text)
                                else:
                                    st.session_state.limitations_future_research = text
                                    st.success("Limitation of the study and suggestions for future research generated successfully!")
                        else:
                            st.error("No non-demographic tables found for limitations generation.")

            with col4:
                if st.button("🚀 Generate Full Report Now", use_container_width=True):
                    if not api_key:
                        st.error("Gemini API key is missing.")
                    else:
                        non_demo = [
                            t for t in st.session_state.interpreted_tables
                            if t["section_tag"] != "DEMOGRAPHIC"
                        ]
                        if non_demo:
                            with st.spinner("Generating comprehensive report..."):
                                comprehensive_prompt = build_comprehensive_report_prompt(
                                    non_demo,
                                    st.session_state.empirical_review_text,
                                    st.session_state.study_context_text,
                                )
                                comprehensive_text = gemini_text(api_key, comprehensive_prompt, model_name)

                                if comprehensive_text.startswith("ERROR:"):
                                    st.error(comprehensive_text)
                                else:
                                    st.session_state.comprehensive_report = comprehensive_text
                                    st.success("Comprehensive report generated successfully!")
                        else:
                            st.error("No non-demographic tables found for full report generation.")

        if st.session_state.comprehensive_report:
            st.markdown("### COMPREHENSIVE REPORT")
            st.markdown(st.session_state.comprehensive_report)
        else:
            if st.session_state.discussion_of_findings:
                st.markdown("### DISCUSSION OF FINDINGS")
                st.markdown(st.session_state.discussion_of_findings)

            if st.session_state.final_outputs:
                st.markdown("### CONCLUSION, IMPLICATION OF THE STUDY, AND RECOMMENDATIONS")
                st.markdown(st.session_state.final_outputs)

            if st.session_state.limitations_future_research:
                st.markdown("### LIMITATION OF THE STUDY AND SUGGESTIONS FOR FUTURE RESEARCH")
                st.markdown(st.session_state.limitations_future_research)

        if st.session_state.interpreted_tables:
            st.markdown("### DOWNLOAD CURRENT WORK")

            current_full_content = build_complete_report_content()

            if st.button("📄 Prepare Full Report (.docx)", use_container_width=True):
                if not current_full_content.strip():
                    st.warning("There is no generated content to download yet.")
                else:
                    st.session_state.docx_buffer = create_text_download_docx(current_full_content)
                    st.success("DOCX report prepared successfully!")

            if st.session_state.docx_buffer is not None:
                st.download_button(
                    "📥 Download Full Report (.docx)",
                    data=st.session_state.docx_buffer.getvalue(),
                    file_name="abt_data_analyst_full_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    st.markdown("</div>", unsafe_allow_html=True)


# =========================================================
# TAB 4 - HISTORY
# =========================================================
with tab4:
    st.markdown("<div class='section-card'>", unsafe_allow_html=True)
    st.subheader("Saved Histories")

    st.markdown("### FULL REPORTS")

    full_reports_content = build_full_reports_content()

    if full_reports_content.strip():
        full_reports_docx = create_text_download_docx(full_reports_content)
        st.download_button(
            "📥 Full Reports (.docx)",
            data=full_reports_docx.getvalue(),
            file_name="abt_data_analyst_full_reports.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_full_reports_docx",
        )
    else:
        st.info("No generated report sections are available yet for Full Reports download.")

    if (
        st.session_state.interpreted_tables
        or st.session_state.summary_of_findings
        or st.session_state.discussion_of_findings
        or st.session_state.final_outputs
        or st.session_state.limitations_future_research
    ):
        st.markdown("### SAVE CURRENT WORKS")

        history_name = st.text_input("Enter a name for this history:", key="history_name_input")

        if st.button("💾 Save Current Works to History", use_container_width=True):
            if not history_name.strip():
                st.warning("Please enter a name for the history.")
            else:
                save_content = build_full_reports_content()

                if save_content.strip():
                    save_history(
                        history_name.strip(),
                        "full_report",
                        save_content
                    )
                    st.success(f"Saved '{history_name.strip()}' to history!")
                    st.session_state.history_name_input = ""
                else:
                    st.error("No content available to save. Please generate some report sections first.")
    else:
        st.info("No current works available to save. Please generate some report sections first.")

    histories = load_history(200)

    if not histories:
        st.info("No saved history yet.")
    else:
        st.markdown(
            """
        <div class='info-box'>
        You can open any saved history, leave it as it is, delete it whenever you want,
        or download it again in .docx format.
        </div>
        """,
            unsafe_allow_html=True,
        )

        for item in histories:
            st.markdown("<div class='history-card'>", unsafe_allow_html=True)
            st.markdown(f"**{item['title']}**")
            st.caption(f"{item['created_at']} • {item['content_type']}")

            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("📂 Open", key=f"open_{item['id']}", use_container_width=True):
                    st.session_state.opened_history_title = item["title"]
                    st.session_state.opened_history_content = item["content"]

            with col2:
                docx_data = create_text_download_docx(item["content"])
                st.download_button(
                    "📥 Download (.docx)",
                    data=docx_data.getvalue(),
                    file_name=f"{item['title'].replace(' ', '_').lower()}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{item['id']}",
                    use_container_width=True,
                )

            with col3:
                if st.button("🗑 Delete", key=f"delete_{item['id']}", use_container_width=True):
                    delete_history(item["id"])
                    if st.session_state.opened_history_title == item["title"]:
                        st.session_state.opened_history_title = ""
                        st.session_state.opened_history_content = ""
                    st.rerun()

            with st.expander("Quick Preview"):
                st.text_area(
                    f"preview_{item['id']}",
                    item["content"],
                    height=220,
                    disabled=True,
                    label_visibility="collapsed",
                )

            st.markdown("</div>", unsafe_allow_html=True)

    if st.session_state.opened_history_content:
        st.markdown("### OPENED HISTORY")
        st.markdown(f"**{st.session_state.opened_history_title}**")
        st.text_area(
            "Opened History Content",
            st.session_state.opened_history_content,
            height=400,
            disabled=True,
        )

        opened_docx = create_text_download_docx(st.session_state.opened_history_content)
        st.download_button(
            "📥 Download Opened History (.docx)",
            data=opened_docx.getvalue(),
            file_name=f"{st.session_state.opened_history_title.replace(' ', '_').lower()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_opened_history_docx",
        )

    st.markdown("</div>", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)
